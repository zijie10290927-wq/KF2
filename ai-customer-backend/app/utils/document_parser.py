"""多格式文档解析器：PDF / Word / TXT / Markdown。

大文件解析必须在后台任务中执行，严禁在 API 请求线程中同步解析。
"""

import io
import logging
from typing import BinaryIO

logger = logging.getLogger(__name__)


def parse_pdf(file_obj: BinaryIO) -> str:
    """PyMuPDF (fitz) 逐页解析 PDF，去除页码噪声。"""
    try:
        import fitz  # PyMuPDF
    except ImportError as e:  # pragma: no cover
        raise ImportError("PyMuPDF 未安装，请执行 `pip install pymupdf`") from e

    text_parts: list[str] = []
    # 兼容 bytes / file-like
    if isinstance(file_obj, (bytes, bytearray)):
        doc = fitz.open(stream=bytes(file_obj), filetype="pdf")
    else:
        data = file_obj.read()
        file_obj.seek(0)
        doc = fitz.open(stream=data, filetype="pdf")

    try:
        for page in doc:
            page_text = page.get_text("text")
            # 去除页脚页码（单行纯数字）
            cleaned_lines = [
                line for line in page_text.splitlines()
                if not line.strip().isdigit()
            ]
            text_parts.append("\n".join(cleaned_lines))
    finally:
        doc.close()
    return "\n\n".join(text_parts).strip()


def parse_docx(file_obj: BinaryIO) -> str:
    """python-docx 遍历 paragraphs，join with \\n。"""
    try:
        from docx import Document
    except ImportError as e:  # pragma: no cover
        raise ImportError("python-docx 未安装，请执行 `pip install python-docx`") from e

    if isinstance(file_obj, (bytes, bytearray)):
        doc = Document(io.BytesIO(bytes(file_obj)))
    else:
        data = file_obj.read()
        file_obj.seek(0)
        doc = Document(io.BytesIO(data))

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()


def parse_txt(file_obj: BinaryIO) -> str:
    """直接读取文本，utf-8 优先，兜底 GBK。"""
    if isinstance(file_obj, (bytes, bytearray)):
        raw = bytes(file_obj)
    else:
        raw = file_obj.read()
        file_obj.seek(0)
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return raw.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore").strip()


def parse_md(file_obj: BinaryIO) -> str:
    """直接读取 markdown 原文。"""
    return parse_txt(file_obj)


def parse_auto(file_obj: BinaryIO, filename: str) -> str:
    """按扩展名自动路由解析器。

    Args:
        file_obj: 二进制文件对象或 bytes
        filename: 文件名（用于判断扩展名）

    Returns:
        解析后的纯文本
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parsers = {
        "pdf": parse_pdf,
        "docx": parse_docx,
        "doc": parse_docx,
        "txt": parse_txt,
        "md": parse_md,
        "markdown": parse_md,
    }
    parser = parsers.get(ext)
    if parser is None:
        logger.warning("Unsupported file type: %s, fallback to txt", ext)
        parser = parse_txt
    return parser(file_obj)
